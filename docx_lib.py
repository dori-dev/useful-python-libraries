"""docx library
"""
from docx import Document
from docx.shared import Inches
# pip install python-docx


class Introduction:
    """Introduction Object
    """

    def __init__(self, name):
        self.document = Document()
        self.create_header("Mohammad Dori")
        text = [
            'I am Developer and I am here for help you!',
            'I learning advanced of js, python and django :)',
            'this text for right align',
            'this text for justified'
        ]
        self.add_text(text)
        self.document.add_page_break()
        self.document.add_paragraph('hello :) now i am in page two :/\n')
        self.document.add_picture('assets/img.jpg', width=Inches(4))
        self.add_section('python', 'js')
        data = [
            ['language', 'python', 'javascript', 'cpp'],
            ['year', '1991', '1995', '1985'],
            ['power', 'AI', 'Web', 'OS'],
            ['love', '100%', '100%', '80%']
        ]
        style = 'Medium Shading 1 Accent 2'
        self.add_table(data, style)
        self.document.save(name)

    def create_header(self, msg: str):
        """create header of doc

        Args:
            msg (str): msg of header
        """
        # level form 0 to 9
        self.document.add_heading(msg, level=0)

    def add_text(self, text: list):
        """add text to page

        Args:
            text (list): list of text line
        """
        # alignments => 0=left, 1=center, 2=right, 3=justified
        for number, line in enumerate(text):
            self.document.add_paragraph(line).alignment = number

    def add_section(self, head_text: str, foot_text: str):
        """add section header and footer of all pages

        Args:
            head_text (str): header text
            foot_text (str): footer text
        """
        section_file = self.document.sections
        header = section_file[0].header.paragraphs[0]
        footer = section_file[0].footer.paragraphs[0]
        header.text = head_text
        header.alignment = 1
        footer.text = foot_text
        footer.alignment = 1

    def add_table(self, data: list, style: str):
        """add table with `data`

        Args:
            data (list): data of table
            style (str): style of table
        """
        table = self.document.add_table(rows=4, cols=4, style=None)
        table.style = style
        for i in range(4):
            row = table.rows[i]
            for j in range(4):
                cell = row.cells[j]
                cell.text = data[i][j]


if __name__ == '__main__':
    for num in range(1, 4):
        file_name = f'document-{str(num)}.docx'
        Introduction(file_name)
